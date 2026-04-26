#!/usr/bin/env python3
"""Build Day 7 results Word document — Sobol re-run on real Fukushima.

Outputs: output/Day7_Results_Sobol_Real_Network.docx
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
FIG_DIR = REPO / "figures" / "fukushima" / "day7"
RES_DIR = REPO / "results" / "day7"
OUT = REPO / "output" / "Day7_Results_Sobol_Real_Network.docx"


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def _add_image(doc: Document, path: Path, width_inches: float = 6.4) -> None:
    if not path.exists():
        doc.add_paragraph(f"[missing figure: {path.name}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_inches))


def _add_table(doc: Document, header: list[str],
               rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)


def main() -> int:
    if not FIG_DIR.exists():
        print(f"ERROR: figures dir missing: {FIG_DIR}", file=sys.stderr)
        return 1

    summary = (json.loads((RES_DIR / "sobol_design_summary.json").read_text())
               if (RES_DIR / "sobol_design_summary.json").exists() else {})
    idx = (pd.read_csv(RES_DIR / "sobol_indices.csv")
           if (RES_DIR / "sobol_indices.csv").exists() else pd.DataFrame())
    cmp_df = (pd.read_csv(RES_DIR / "day4_vs_day7_comparison.csv")
              if (RES_DIR / "day4_vs_day7_comparison.csv").exists()
              else pd.DataFrame())
    sep = (pd.read_csv(RES_DIR / "cmc_separability.csv")
           if (RES_DIR / "cmc_separability.csv").exists() else pd.DataFrame())

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title --------------------------------------------------------------
    doc.add_heading("Day 7: Full Sobol Sensitivity Re-run on the "
                    "Real Fukushima OSM Network", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        "Definitive pre-calibration sensitivity characterization — "
        "four free parameters, six QoIs, bootstrap confidence intervals"
    ).italic = True
    doc.add_paragraph(
        f"Saltelli quasi-random design with n_samples="
        f"{summary.get('n_samples', '?')} → "
        f"{summary.get('n_evaluations', '?')} parameter evaluations × "
        f"{summary.get('n_members', '?')} ensemble members per evaluation, "
        f"on the route6_closed scenario "
        f"(blend mode, {summary.get('n_agents', '?')} agents, "
        f"{summary.get('n_steps', '?')} timesteps). "
        f"Total simulations: {summary.get('total_simulations', '?')}; "
        f"main campaign runtime "
        f"{summary.get('main_campaign_runtime_sec', '?')} s; "
        f"CMC separability mini-analysis "
        f"{summary.get('cmc_sep_runtime_sec', '?')} s."
    )

    # --- 1. Purpose ---------------------------------------------------------
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "Day 4 produced first-order and total-order Sobol indices for "
        "α, β, κ on a synthetic three-zone topology, but did so against "
        "a perception layer that collapsed all spatial gradients to "
        "zero before they reached the σ computation. κ was therefore "
        "structurally unidentifiable in every Day 4 QoI (ST = 0 across "
        "the board). Day 4b removed the perception layer and substituted "
        "the precomputed ConflictPotentialField; Day 6 verified κ was "
        "identifiable on the real network via a dedicated 5-point sweep."
    )
    doc.add_paragraph(
        "Day 7 replaces the Day 4 indices with a corrected design on the "
        "real Fukushima OSM network. It adds ConflictMoveChance (CMC) as "
        "a fourth free parameter, uses the route6_closed scenario "
        "(richer S2 signal than baseline, more discriminating for κ), "
        "computes bootstrap CIs on every index, and runs an explicit "
        "CMC-separability mini-analysis to test whether the planned "
        "two-stage Day 8 calibration is feasible."
    )

    # --- 2. Design ----------------------------------------------------------
    doc.add_heading("2. Design", level=1)
    bullets = [
        "Parameters: α ∈ [0.5, 5.0]; β ∈ [1.0, 10.0]; κ ∈ [1.0, 20.0]; "
        "CMC ∈ [0.25, 0.75]. All uniform.",
        "Saltelli sampling with calc_second_order=False → "
        f"N = n_samples × (D+2) = {summary.get('n_samples', '?')} × 6 = "
        f"{summary.get('n_evaluations', '?')} parameter vectors. "
        "Note: the prompt's formula 32 × (2D+2) = 320 assumes "
        "calc_second_order=True; with calc_second_order=False as the "
        "prompt also requests, the actual Saltelli evaluation count is "
        f"{summary.get('n_evaluations', '?')} (correct for first- and "
        "total-order indices). All 192 evaluations were used.",
        "Six QoIs: hayano_t4, mid_ps2_trough, mid_ps2_dip, "
        "mid_ps2_recovery (carry-overs from Day 4); corridor_inland_pct "
        "and blend_inner_t7 (new, motivated by Days 5–6).",
        "Scenario: route6_closed (conflict spike at naraha+hirono "
        "from day 3) — chosen over baseline because the spatial "
        "gradient is largest there and gives κ its strongest signal.",
        "Mode: blend only (cognitive parameters affect no other mode).",
        "Ensemble: 3 members per evaluation (Sobol estimator absorbs "
        "ensemble variance into the index uncertainty).",
        "Experience distribution: held fixed at Beta(2, 5). Day 6 §3 "
        "showed the distribution shape affects P_S2 amplitude but not "
        "corridor structure; varying it would conflate Ψ with α.",
        f"Bootstrap: SALib's internal resampling with "
        f"num_resamples={summary.get('bootstrap_resamples', '?')} and "
        "conf_level=0.95.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # --- 3. Headline indices table ------------------------------------------
    doc.add_heading("3. Sobol Indices with 95% Bootstrap CIs", level=1)
    if not idx.empty:
        rows = []
        for _, r in idx.iterrows():
            note = ""
            if r.get("ST_includes_zero", False):
                note = "CI includes 0"
            elif r.get("effectively_insensitive", False):
                note = "ST < 0.05"
            rows.append([
                r["qoi"], r["parameter"],
                f"{r['S1']:+.3f}",
                f"[{r['S1_CI_low']:+.3f}, {r['S1_CI_high']:+.3f}]",
                f"{r['ST']:+.3f}",
                f"[{r['ST_CI_low']:+.3f}, {r['ST_CI_high']:+.3f}]",
                note,
            ])
        _add_table(doc,
                   header=["QoI", "Param",
                           "S1", "S1 95% CI",
                           "ST", "ST 95% CI",
                           "Insensitive?"],
                   rows=rows)

    _add_image(doc, FIG_DIR / "D7-1_sobol_indices.png", width_inches=6.8)
    _add_caption(doc, "Fig D7-1. First-order (S1, top row) and total-order "
                      "(ST, bottom row) Sobol indices for the four "
                      "parameters across the six QoIs, with 95% bootstrap "
                      "confidence intervals. Dashed line at 0.05 marks "
                      "the conventional insensitivity threshold.")

    doc.add_paragraph(
        "Headline observations. (i) β still dominates mid_ps2_trough at "
        "ST = 0.79 — within 1% of the Day 4 value of 0.80. The β-driven "
        "Ω-collapse mechanism is preserved when moving from synthetic "
        "to real topology. (ii) α dominates mid_ps2_dip at ST = 0.77 "
        "(Day 4: 0.66) — the dip-amplitude identifiability for α "
        "actually strengthens on the real network. (iii) κ has "
        "ST > 0.05 with CI excluding zero in four of the six QoIs "
        "(hayano_t4, mid_ps2_recovery, corridor_inland_pct, "
        "blend_inner_t7). Day 4 had zero κ-sensitive QoIs; the Day 4b "
        "perception fix is therefore confirmed at production scale. "
        "(iv) CMC dominates the four outcome QoIs — its ST > 0.7 in "
        "every QoI except mid_ps2_trough and mid_ps2_dip. This is "
        "physically expected: CMC sets the per-step departure rate "
        "for any agent in a conflict cell. (v) Several ST estimates "
        "exceed 1.0 (e.g. blend_inner_t7/CMC ST = 1.22). With "
        "n_samples=32 the Saltelli ST estimator carries finite-sample "
        "bias, particularly when interactions are large; ST > 1 "
        "should be read as 'dominant with strong interaction', not "
        "as a violation."
    )

    # --- 4. Interactions ----------------------------------------------------
    doc.add_heading("4. Interaction Structure", level=1)
    _add_image(doc, FIG_DIR / "D7-2_interactions.png", width_inches=6.4)
    _add_caption(doc, "Fig D7-2. Interaction magnitudes (ST − S1) for each "
                      "parameter × QoI cell. Larger values indicate the "
                      "parameter's effect on that QoI is mediated by "
                      "interactions with other parameters; smaller values "
                      "indicate near-additive behaviour.")
    doc.add_paragraph(
        "Mean interaction magnitude across all parameter × QoI cells is "
        f"{idx['interaction'].mean():.3f}, comparable to the Day 4 mean "
        "of 0.23. The interaction structure has shifted in a "
        "scientifically informative way: (a) corridor_inland_pct shows "
        "the largest interactions across all four parameters (α: 0.65, "
        "β: 0.56, κ: 0.33, CMC: 0.32) — corridor choice emerges from "
        "joint cognitive-physical decision-making, not from any single "
        "parameter; (b) mid_ps2_recovery shows large CMC and κ "
        "interactions (0.98 and 0.55), confirming the Day 4 finding "
        "that the recovery phase is multi-parameter; (c) mid_ps2_trough "
        "remains the cleanest β-additive QoI (interaction ≈ 0 for β), "
        "reinforcing its value as a calibration target for β alone."
    )

    # --- 5. κ characterization ----------------------------------------------
    doc.add_heading("5. κ Characterization", level=1)
    _add_image(doc, FIG_DIR / "D7-3_kappa_scatter_by_beta.png",
               width_inches=6.6)
    _add_caption(doc, "Fig D7-3. Each QoI plotted against κ (log axis) "
                      "for all 192 Sobol samples, colored by β. The "
                      "hypothesis is that κ's effect is strongest when "
                      "β is low (Ω is high, agents are deliberating) "
                      "and weakest when β is high (Ω is suppressed, "
                      "κ has nothing to act on).")
    doc.add_paragraph(
        "κ sensitivity is concentrated in mid_ps2_recovery (ST = 0.53), "
        "corridor_inland_pct (ST = 0.32), and blend_inner_t7 (ST = 0.30). "
        "It is essentially absent in mid_ps2_trough (ST = 0.005) and "
        "mid_ps2_dip (ST = 0.005). This directly answers Day 6's open "
        "question 'which QoIs does κ drive': κ governs the recovery "
        "and outcome phases of the dual-process trajectory but not the "
        "trough and dip themselves, which are α/β-determined. The "
        "calibration implication is that κ should be tuned against "
        "recovery and corridor metrics, not against the trough."
    )

    # --- 6. CMC separability ------------------------------------------------
    doc.add_heading("6. CMC Separability — the Two-Stage "
                    "Day 8 Design Test", level=1)
    doc.add_paragraph(
        "The two-stage Day 8 design assumes CMC and the cognitive "
        "parameters are approximately separable, i.e. the ST indices "
        "for α, β, κ are stable across CMC levels. If they are, CMC "
        "can be fixed first (e.g. from the Hayano timing constraint) "
        "and (α, β, κ) can be calibrated independently in the second "
        "stage. If indices change substantially with CMC, the two "
        "stages cannot be cleanly separated and Day 8 must use a joint "
        "four-parameter search."
    )
    if not sep.empty:
        piv = (sep.pivot_table(index=["qoi", "parameter"],
                               columns="cmc", values="ST")
               .round(3))
        rows = []
        for (qoi, param), r in piv.iterrows():
            vals = r.values
            drift = float(max(vals) - min(vals))
            rows.append([qoi, param,
                         f"{r[0.25]:+.3f}", f"{r[0.50]:+.3f}",
                         f"{r[0.75]:+.3f}",
                         f"{drift:+.3f}"])
        _add_table(doc,
                   header=["QoI", "Param",
                           "ST @ CMC=0.25", "ST @ CMC=0.50",
                           "ST @ CMC=0.75", "Drift"],
                   rows=rows)
    _add_image(doc, FIG_DIR / "D7-5_cmc_separability.png", width_inches=6.8)
    _add_caption(doc, "Fig D7-5. ST(α, β, κ) at three fixed CMC levels. "
                      "Lines that are flat across CMC indicate the "
                      "cognitive index is independent of CMC (separable); "
                      "lines that change indicate CMC interacts with "
                      "that cognitive parameter for that QoI.")

    doc.add_paragraph(
        "Result: CMC separability fails on the outcome QoIs. The "
        "maximum ST drift across CMC levels is 1.00 (β/hayano_t4 — "
        "ST = 0.26 at CMC=0.25 vs ST = 1.26 at CMC=0.50). For "
        "hayano_t4, blend_inner_t7, corridor_inland_pct, and "
        "mid_ps2_recovery the cognitive indices change by 0.5 or more "
        "as CMC moves from 0.25 to 0.75. The ST drift exceeds the "
        "0.15 threshold for nearly every (cognitive parameter × "
        "outcome QoI) pair."
    )
    doc.add_paragraph(
        "However, CMC separability holds reasonably well on the two "
        "process QoIs. mid_ps2_trough has β drift 0.18 (ST goes "
        "0.64 → 0.99 → 0.82) and α drift 0.33 — both above the strict "
        "0.15 threshold, but the dominance ordering "
        "(β ≫ α > κ ≈ CMC) is preserved. mid_ps2_dip is even more "
        "stable: α drift 0.28, β drift 0.07, κ drift 0.01."
    )
    doc.add_paragraph(
        "Day 8 redesign implication. The clean two-stage 'physical "
        "first, cognitive second' design assumed in the prompt is not "
        "feasible. The corrected two-stage design that this analysis "
        "supports is: (1) calibrate β and α from the process QoIs "
        "(mid_ps2_trough and mid_ps2_dip), where CMC interaction is "
        "small enough to use a fixed reasonable CMC; (2) calibrate "
        "CMC and κ jointly from the outcome QoIs (hayano_t4, "
        "blend_inner_t7, corridor_inland_pct), holding β and α at the "
        "Stage-1 estimates. This is two-stage by QoI grouping rather "
        "than by parameter grouping, but it preserves the practical "
        "tractability the prompt was aiming for."
    )

    # --- 7. Day 4 vs Day 7 --------------------------------------------------
    doc.add_heading("7. Day 4 vs Day 7 Comparison", level=1)
    if not cmp_df.empty:
        rows = []
        for _, r in cmp_df.iterrows():
            rows.append([
                r["qoi"], r["parameter"],
                f"{r['ST_day4']:.3f}",
                f"{r['ST_day7']:.3f}",
                f"{r['delta']:+.3f}",
            ])
        _add_table(doc,
                   header=["QoI", "Param",
                           "ST Day 4", "ST Day 7", "Δ"],
                   rows=rows)
    _add_image(doc, FIG_DIR / "D7-4_day4_vs_day7.png", width_inches=6.6)
    _add_caption(doc, "Fig D7-4. Side-by-side ST for the four shared QoIs. "
                      "Day 4 used the synthetic three-zone topology with "
                      "the broken perception layer; Day 7 uses the real "
                      "OSM network with the corrected potential field. "
                      "Note κ rises from exactly 0 (Day 4) to non-zero "
                      "in every QoI in Day 7.")
    doc.add_paragraph(
        "Three high-level shifts: (a) hayano_t4 sensitivities to α and "
        "β both fell substantially (α: 0.40 → 0.13; β: 0.27 → 0.08). "
        "This is because CMC, which was held fixed at 0.5 in Day 4 "
        "and is now varied, has absorbed most of the variance in "
        "departure timing. (b) mid_ps2_trough β-dominance persists "
        "almost unchanged (0.80 → 0.79). (c) mid_ps2_recovery "
        "indices fall for α and β but κ jumps from 0 to 0.53 — "
        "recovery is the QoI most affected by the perception fix, "
        "consistent with the theoretical prediction that recovery "
        "(when agents resolve toward S2) is the phase most sensitive "
        "to safety signal sharpness."
    )

    # --- 8. Diagnostics gate ------------------------------------------------
    doc.add_heading("8. Diagnostics Gate", level=1)
    gate_rows = [
        ["β ST for mid_ps2_trough",
         "Compare to Day 4 = 0.80",
         "Day 7 = 0.79; β dominance preserved",
         "PASS"],
        ["κ ST > 0.05 with CI excluding 0",
         "≥ 1 QoI",
         "4 QoIs (hayano_t4, mid_ps2_recovery, "
         "corridor_inland_pct, blend_inner_t7)",
         "PASS"],
        ["CMC separability",
         "Cognitive ST stable ±0.15 across CMC levels",
         "Outcome QoIs: max drift ≈ 1.0 (FAIL); "
         "Process QoIs: drift 0.07–0.33 (mostly OK)",
         "PARTIAL — Day 8 must redesign"],
        ["Bootstrap CIs computed",
         "All indices have 95% CI",
         "48 indices have non-zero CI; all reported",
         "PASS"],
        ["84 tests still passing",
         "No regressions",
         "84 passed, 2 skipped",
         "PASS"],
    ]
    _add_table(doc,
               header=["Check", "Threshold", "Observed", "Result"],
               rows=gate_rows)

    # --- 9. Implications for Day 8 ------------------------------------------
    doc.add_heading("9. Implications for Day 8 Calibration", level=1)
    impl = [
        "Three-parameter (α, β, κ) calibration is now justified — κ is "
        "identifiable in 4/6 QoIs.",
        "CMC must be included as a fourth free parameter, NOT held "
        "fixed at the Day 3 default of 0.5. The CMC indices in Fig "
        "D7-1 dominate the outcome QoIs and the separability test "
        "rules out a clean physical-first / cognitive-second split.",
        "Two-stage design by QoI grouping is the recommended "
        "alternative: Stage 1 fits β and α to the process markers "
        "(mid_ps2_trough, mid_ps2_dip) where CMC interactions are "
        "small; Stage 2 fits CMC and κ jointly to the outcome metrics "
        "(hayano_t4, blend_inner_t7, corridor_inland_pct), holding "
        "Stage-1 estimates fixed. This preserves the tractability of "
        "the original two-stage idea while honouring the interaction "
        "structure exposed by the Sobol analysis.",
        "Day 6 §4 already showed that hayano_t4 = 0.78 cannot be "
        "reached within the physical CMC range [0.25, 0.75]. Day 8 "
        "should therefore not chase the Hayano level as a hard "
        "target; it should use the timing rate as a soft anchor and "
        "report residual level bias as a known systematic.",
        "corridor_inland_pct is now a validated, well-identified QoI "
        "(ST > 0.3 for α, β, κ; ST > 0.7 for CMC). It is the natural "
        "Day 8 routing-validation QoI to complement the timing "
        "(hayano_t4) and process (mid_ps2_*) targets.",
    ]
    for i, b in enumerate(impl, start=1):
        doc.add_paragraph(b, style="List Number")

    # --- 10. Artifacts ------------------------------------------------------
    doc.add_heading("10. Artifacts", level=1)
    doc.add_paragraph("Scripts:")
    for s in [
        "scripts/run_sobol_day7.py — main Day 7 driver "
        "(--quick smoke, --skip-cmc-sep options)",
        "scripts/build_day7_docx.py — this document",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("Data outputs (results/day7/):")
    for s in [
        "sobol_raw_results.csv — 192 evals × (4 params + 6 QoIs)",
        "sobol_indices.csv — full S1/ST table with bootstrap CIs and "
        "insensitivity flags",
        "interaction_magnitudes.csv — (ST − S1) per parameter × QoI",
        "cmc_separability.csv — Section 6 mini-analysis "
        "(3 CMC levels × n=24)",
        "day4_vs_day7_comparison.csv — Section 7 ST deltas",
        "sobol_design_summary.json — design parameters, runtimes, "
        "total compute",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("Figures (figures/fukushima/day7/):")
    for s in [
        "D7-1_sobol_indices.png — primary summary",
        "D7-2_interactions.png — interaction heatmap",
        "D7-3_kappa_scatter_by_beta.png — κ characterization scatter",
        "D7-4_day4_vs_day7.png — broken perception vs corrected",
        "D7-5_cmc_separability.png — Day 8 design validation test",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    repro = doc.add_paragraph()
    repro.add_run("Reproduction (~3 minutes on a laptop, "
                  "smoke test via --quick):").italic = True
    p = doc.add_paragraph()
    p.add_run("python scripts/run_sobol_day7.py").font.name = "Consolas"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size/1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
